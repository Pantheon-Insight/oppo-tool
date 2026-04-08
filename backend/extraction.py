import pdfplumber
from docx import Document
import json
import os
from typing import List, Tuple
import anthropic
from models import Attack, Severity


def extract_text_from_pdf(pdf_path: str) -> Tuple[str, int]:
    """
    Extract text from PDF file.
    Returns tuple of (full_text, page_count).
    """
    full_text = ""
    page_count = 0

    try:
        with pdfplumber.open(pdf_path) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    full_text += text + "\n\n"
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        raise

    return full_text, page_count


def extract_text_from_docx(docx_path: str) -> Tuple[str, int]:
    """
    Extract text from DOCX file.
    Returns tuple of (full_text, paragraph_count_as_page_estimate).
    """
    try:
        doc = Document(docx_path)
        full_text = "\n".join([para.text for para in doc.paragraphs])
        page_count = len(doc.paragraphs)  # Estimate pages as paragraph count
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        raise

    return full_text, page_count


def chunk_text(text: str, chunk_size: int = 8000, overlap: int = 500) -> List[str]:
    """
    Split text into overlapping chunks for Claude processing.
    Uses character-based chunking with overlap to preserve context.
    """
    chunks = []
    start = 0

    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunk = text[start:end]
        chunks.append(chunk)

        # Move start for next chunk with overlap
        start = end - overlap

    return chunks


def extract_attacks_from_chunk(chunk_text: str, chunk_number: int) -> List[dict]:
    """
    Use Claude API to extract attacks from a single chunk.
    Returns list of attack dictionaries.
    """
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise ValueError("ANTHROPIC_API_KEY environment variable not set")

    client = anthropic.Anthropic(api_key=api_key)
    model = os.getenv("CLAUDE_MODEL", "claude-sonnet-4-20250514")

    system_prompt = """You are an opposition research extraction specialist. Your task is to extract EVERY possible attack, hit, negative fact, or damaging information from the provided text.

An "attack" includes but is not limited to:
- Voting record and specific votes taken
- Campaign contributions and donations (sources and amounts)
- Endorsements received or given
- Political associations and relationships
- Caucus memberships and group affiliations
- Professional history and employment records
- Financial disclosures and conflicts of interest
- Direct quotes (especially controversial or hypocritical ones)
- Personal facts (age, residency, family background)
- Legal issues, lawsuits, or investigations
- Hypocrisy angles (contradictions between positions and actions)
- Business dealings and corporate affiliations
- Award rejections or controversies
- Statements on policy issues (especially flip-flops)
- Campaign violations or ethics complaints
- Social media controversies
- Association with controversial figures or organizations
- Geographic considerations (where they live, represent, or have connections)
- Any factual information that could be used in negative campaigning

For EACH attack found:
1. Assign a category (e.g., "Voting Record", "Campaign Finance", "Personal", "Professional", "Statements", "Associations", "Financial", "Legal", etc.)
2. Write the attack as a concise, damaging statement
3. Include key details/evidence that supports the attack
4. Assign severity: "Major" (high impact), "Moderate" (significant), "Minor" (small but notable), or "Niche" (specialized audiences)

Return results as a JSON array of objects with this exact structure:
{
  "attacks": [
    {
      "category": "string",
      "attack": "string",
      "key_detail": "string",
      "severity": "Major|Moderate|Minor|Niche"
    }
  ]
}

Extract EVERYTHING. Be comprehensive. Err on the side of including more rather than fewer attacks."""

    try:
        response = client.messages.create(
            model=model,
            max_tokens=4000,
            system=system_prompt,
            messages=[
                {
                    "role": "user",
                    "content": f"Extract all attacks from this text chunk #{chunk_number}:\n\n{chunk_text}"
                }
            ]
        )

        response_text = response.content[0].text

        # Parse JSON response
        try:
            # Try to extract JSON from response
            if "```json" in response_text:
                json_str = response_text.split("```json")[1].split("```")[0].strip()
            elif "```" in response_text:
                json_str = response_text.split("```")[1].split("```")[0].strip()
            else:
                json_str = response_text

            data = json.loads(json_str)
            return data.get("attacks", [])
        except json.JSONDecodeError:
            print(f"Warning: Failed to parse JSON from chunk {chunk_number}")
            return []

    except Exception as e:
        print(f"Error extracting attacks from chunk {chunk_number}: {e}")
        raise


def deduplicate_attacks(attacks: List[dict]) -> List[dict]:
    """
    Remove duplicate attacks across chunks.
    Uses attack title + key_detail as deduplication key.
    """
    seen = set()
    deduped = []

    for attack in attacks:
        # Create a normalized key for comparison
        key = (
            attack.get("attack", "").lower().strip(),
            attack.get("key_detail", "").lower().strip(),
            attack.get("category", "").lower().strip()
        )

        if key not in seen:
            seen.add(key)
            deduped.append(attack)

    return deduped


def extract_all_attacks(text: str, filename: str) -> List[Attack]:
    """
    Main extraction pipeline:
    1. Split text into chunks
    2. Extract attacks from each chunk
    3. Deduplicate
    4. Convert to Attack models with sequential numbering
    """
    # Split text into chunks
    chunks = chunk_text(text)
    print(f"Processing {len(chunks)} chunks from {filename}")

    all_attacks = []

    # Extract from each chunk
    for i, chunk in enumerate(chunks):
        print(f"Extracting from chunk {i+1}/{len(chunks)}...")
        try:
            chunk_attacks = extract_attacks_from_chunk(chunk, i + 1)
            all_attacks.extend(chunk_attacks)
        except Exception as e:
            print(f"Error processing chunk {i+1}: {e}")
            continue

    # Deduplicate
    deduped_attacks = deduplicate_attacks(all_attacks)
    print(f"Found {len(all_attacks)} attacks, {len(deduped_attacks)} after deduplication")

    # Convert to Attack models with numbering
    result = []
    for i, attack_data in enumerate(deduped_attacks, 1):
        try:
            severity_str = attack_data.get("severity", "Minor").lower()
            if severity_str == "major":
                severity = Severity.MAJOR
            elif severity_str == "moderate":
                severity = Severity.MODERATE
            elif severity_str == "minor":
                severity = Severity.MINOR
            else:
                severity = Severity.NICHE

            attack_obj = Attack(
                number=i,
                category=attack_data.get("category", "Other"),
                attack=attack_data.get("attack", ""),
                key_detail=attack_data.get("key_detail", ""),
                severity=severity,
                notes=attack_data.get("notes")
            )
            result.append(attack_obj)
        except Exception as e:
            print(f"Error converting attack to model: {e}")
            continue

    return result
